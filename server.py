#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
日語文獻 OCR+翻譯 Web 介面後端 (非同步輪詢完全體版)
"""

import sys
import io

# 🎯【核心防禦一：強制注入 Python 全域標準輸出 UTF-8 靈魂】
# 徹底打碎雲端精簡 Linux 容器的 Latin-1 編碼限制，防止背景 log 打印中文時崩潰
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

import base64, json, queue, re, threading, time, uuid
from pathlib import Path
from urllib.parse import quote
import requests

try:
    from flask import Flask, request, Response, send_from_directory, jsonify
except ImportError:
    raise SystemExit("请先安装 Flask:  pip install flask")

app = Flask(__name__)

# 🎯【核心防禦二：強制配置 Flask 序列化與 HTTP 標頭為 UTF-8】
# 這是消滅 /api/task_status 輪詢時噴出 latin-1 錯誤的終極解藥！
app.config['JSON_AS_ASCII'] = False
app.config['JSONIFY_MIMETYPE'] = "application/json; charset=utf-8"

# 存儲任務的執行狀態與最終結果
_task_statuses: dict[str, dict] = {}

# 已完成/出错的任务保留 2 小时（供刷新重连），之后清理以释放内存
_TASK_TTL = 7200

def _prune_tasks():
    now = time.time()
    stale = [k for k, v in list(_task_statuses.items())
             if v.get("state") in ("done", "error") and now - v.get("t0", now) > _TASK_TTL]
    for k in stale:
        _task_statuses.pop(k, None)

GAP_MIN_FRAC  = 0.04   
OCR_JOB_URL = "https://paddleocr.aistudio-app.com/api/v2/ocr/jobs"
OCR_MODEL   = "PaddleOCR-VL-1.6"

TRANSLATE_PROMPT = """\
你是资深的{source}→{target}学术翻译。下面是 OCR 得到的一页{source}文献（Markdown 格式，可能含标题、表格、脚注、页眉页脚）。
请翻译成{target}，要求：
1. 保留原有 Markdown 结构（标题层级、表格、列表、强调）。
2. 学术术语、人名、地名、书名尽量准确；没有把握处保留{source}原词并括注译法。
3. 只输出译文本身，不要引用或重复原文，不要任何开场白或说明。
4. 明显 OCR 噪声（断字、个别乱码）可酌情修正，但不臆造内容。
----- 原文开始 -----
{text}
----- 原文结束 -----"""


def _ascii_header(value: str, name: str) -> str:
    """HTTP header 只能放 ASCII；提前拦截混入的中文/全角字符，给出人话报错。"""
    v = (value or "").strip()
    try:
        v.encode("ascii")
    except UnicodeEncodeError:
        raise RuntimeError(
            f"{name} 含非 ASCII 字符（可能复制时混进了中文/全角/多余文字），"
            f"请重新填入纯英文数字的值"
        )
    return v


def _strip_md_noise(text: str) -> str:
    """去掉 OCR/Markdown 里常见的噪声：图片标记、HTML 图片、HTML 注释。"""
    text = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', text)         # Markdown 图片 ![](...)
    text = re.sub(r'<img[^>]*>', '', text)                   # HTML <img ...>
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)  # HTML 注释 <!-- ... -->
    return text


def _parse_ocr_text(text: str) -> list[str]:
    text = _strip_md_noise(text)
    parts = re.split(r'\n\s*[-*]{3,}\s*\n', text)
    pages = [p.strip() for p in parts if p.strip()]
    return pages if pages else [text.strip()]

def _ocr_submit(data: bytes, filename: str, token: str) -> str:
    """提交 OCR 任务，返回 jobId。"""
    token = _ascii_header(token, "PaddleOCR token")
    headers = {"Authorization": f"bearer {token}"}
    
    optional = {
        "useDocOrientationClassify": True,
        "useDocUnwarping": False,
        "useChartRecognition": False,
    }
    
    # 建構 Form 表單時，確保裡面所有 JSON 數據都被純 ASCII 轉義（中文字會變成 \uXXXX）
    form = {
        "model": str(OCR_MODEL), 
        "optionalPayload": json.dumps(optional, ensure_ascii=True)
    }
    
    # 檔案欄位名稱與檔名強行鎖死為純英文，絕不攜帶中文參數流入 requests 庫
    safe_file_field = ("task.pdf", data, "application/pdf")
    files = {"file": safe_file_field}
    
    print(f"[Debug] submitting OCR job to Baidu ...")
    
    # 🎯【終極修正】：將 timeout 直接設為單個 None
    # 這代表：連接時間無限、傳輸檔案數據時間無限、等待響應時間無限！
    # 徹底給跨海網絡上傳留出充裕的時間，杜絕一切 Write Timeout。
    resp = requests.post(OCR_JOB_URL, headers=headers, data=form, files=files, timeout=None)
    
    if resp.status_code != 200:
        raise RuntimeError(f"OCR 提交失败 [{resp.status_code}]: {resp.text[:400]}")
    return resp.json()["data"]["jobId"]

def _ocr_poll(job_id: str, token: str, task_id: str, log_fn=None) -> str:
    token = _ascii_header(token, "PaddleOCR token")
    headers = {"Authorization": f"bearer {token}"}
    while True:
        resp  = requests.get(f"{OCR_JOB_URL}/{job_id}", headers=headers, timeout=30)
        if resp.status_code != 200:
            raise RuntimeError(f"OCR 状态查询失败 [{resp.status_code}]: {resp.text[:400]}")
        data  = resp.json()["data"]
        state = data["state"]
        if state == "done":
            return data["resultUrl"]["jsonUrl"]
        if state == "failed":
            raise RuntimeError(f"OCR 任务失败: {data.get('errorMsg', '未知错误')}")
        if state == "running" and log_fn:
            try:
                prog = data["extractProgress"]
                log_fn(f"OCR 进行中: {prog['extractedPages']}/{prog['totalPages']} 页")
            except (KeyError, TypeError):
                pass
        time.sleep(5)

def _ocr_parse(jsonl_url: str) -> tuple[list[str], list[dict]]:
    resp = requests.get(jsonl_url, timeout=120)
    resp.raise_for_status()
    texts: list[str] = []
    raws:  list[dict] = []
    for line in resp.text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        result = json.loads(line)["result"]
        for res in result.get("layoutParsingResults", []):
            texts.append(res.get("markdown", {}).get("text", ""))
            raws.append(res)
    return texts, raws

def _ocr_file(data: bytes, filename: str, token: str, task_id: str, log_fn=None) -> tuple[list[str], list[dict]]:
    if log_fn:
        log_fn(f"正在上传文件至 OCR 引擎（{len(data)/1048576:.1f} MB）…大文件较慢，请耐心等待，勿刷新")
    job_id = _ocr_submit(data, filename, token)
    if log_fn:
        log_fn(f"任务已提交给云端引擎（ID: {job_id[:8]}...）")
    jsonl_url = _ocr_poll(job_id, token, task_id, log_fn)
    return _ocr_parse(jsonl_url)

def _translate(text: str, target: str, cfg: dict) -> str:
    if not text.strip():
        return "*(本页无文字内容)*"
    provider   = cfg.get("provider", "claude")
    model      = cfg.get("model", "claude-sonnet-4-6")
    key        = _ascii_header(cfg.get("translate_key", ""), "翻译 API key")
    max_tokens = int(cfg.get("max_tokens", 8192))
    source     = cfg.get("source", "日语")   # 源语言；默认日语（本工具首要场景）
    
    if provider == "claude":
        headers = {
            "x-api-key": key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
        body = {
            "model": model,
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": TRANSLATE_PROMPT.format(source=source, target=target, text=text)}],
        }
        resp = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=body, timeout=600)
        if resp.status_code == 200:
            return "".join(b["text"] for b in resp.json()["content"] if b["type"] == "text").strip()
        raise RuntimeError(f"Claude 翻译失败 [{resp.status_code}]: {resp.text[:200]}")
        
    PROVIDER_BASE_URLS = {
        "deepseek": "https://api.deepseek.com/v1",
        "openai":   "https://api.openai.com/v1",
        "gemini":   "https://generativelanguage.googleapis.com/v1beta/openai",
    }
    if provider == "custom":
        base_url = (cfg.get("custom_base_url", "") or "").strip()
        if not base_url:
            raise RuntimeError("自定义服务商需填写 Base URL")
    else:
        base_url = PROVIDER_BASE_URLS.get(provider, "")
        if not base_url:
            raise RuntimeError(f"未知的翻译服务商: {provider}")
    url = base_url.rstrip("/") + "/chat/completions"
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    body = {
        "model": model,
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": TRANSLATE_PROMPT.format(source=source, target=target, text=text)}],
    }
    resp = requests.post(url, headers=headers, json=body, timeout=600)
    if resp.status_code == 200:
        return resp.json()["choices"][0]["message"]["content"].strip()
    raise RuntimeError(f"LLM 翻译失败 [{resp.status_code}]: {resp.text[:200]}")

# ── 背景执行线程 ─────────────────────────────────────────────
def _async_worker(task_id: str, data: bytes, filename: str, cfg: dict):
    st = _task_statuses[task_id]

    def update_log(msg: str, state="running"):
        st["msg"] = msg
        st["state"] = state

    try:
        token    = cfg.get("paddle_token", "")
        source   = cfg.get("source", "日语")
        target   = cfg.get("target", "中文")
        ocr_only = cfg.get("ocr_only", False)
        stem     = cfg.get("stem", "文献")

        ext = Path(filename).suffix.lower()
        update_log("开始处理文献录入...")

        if ext in (".md", ".txt"):
            text = data.decode("utf-8", errors="replace")
            pages = _parse_ocr_text(text)
        else:
            pages, _ = _ocr_file(data, filename, token, task_id, log_fn=update_log)

        n = len(pages)
        # 推送 OCR 结果：前端轮询到 ocr_ready 后即可立即渲染原文 + 书影
        st["ocr_pages"] = pages
        st["n"] = n
        st["ocr_ready"] = True
        st["trans"] = [None] * n

        ocr_md = "\n\n---\n\n".join(f"\n\n{p}" for p in pages)

        if ocr_only:
            st["msg"] = "✓ 仅 OCR 完成"
            st["state"] = "done"
            st["result"] = {"stem": stem, "ocr_pages": pages, "trans_pages": [],
                            "ocr_md": ocr_md, "bilingual_md": ""}
            return

        update_log(f"OCR 解析完毕，共 {n} 页。开始逐页翻译…")

        # 5 路并发翻译；每完成一页就写进 st["trans"][idx]，前端按页增量拉取
        from concurrent.futures import ThreadPoolExecutor, as_completed
        trans_pages = [""] * n

        def _work(idx):
            try:
                return idx, _translate(pages[idx], target, cfg)
            except Exception as e:
                return idx, f"*(第 {idx+1} 頁翻譯出錯: {str(e)})*"

        done = 0
        with ThreadPoolExecutor(max_workers=5) as ex:
            futures = [ex.submit(_work, i) for i in range(n)]
            for fut in as_completed(futures):
                idx, tr = fut.result()
                trans_pages[idx] = tr
                st["trans"][idx] = tr
                done += 1
                st["msg"] = f"已翻譯 {done}/{n} 頁…"

        bi_blocks = []
        for i, (ocr_p, tr_p) in enumerate(zip(pages, trans_pages), 1):
            quoted = "\n".join(f"> {ln}" if ln.strip() else ">" for ln in ocr_p.splitlines())
            bi_blocks.append(f"\n## 第 {i} 页\n\n{quoted}\n\n{tr_p}")
        bilingual_md = f"# {stem} — {source}{target}对照\n\n" + "\n\n---\n\n".join(bi_blocks)

        st["state"] = "done"
        st["msg"] = "✓ 全书解析校读完成！"
        st["result"] = {"stem": stem, "ocr_pages": pages, "trans_pages": trans_pages,
                        "ocr_md": ocr_md, "bilingual_md": bilingual_md}

    except Exception as exc:
        st["state"] = "error"
        st["msg"] = f"错误: {str(exc)}"

# ── 路由端点 ──────────────────────────────────────────────────
@app.route("/api/process", methods=["POST"])
def process():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "no file"}), 400
    
    cfg = json.loads(request.form.get("config", "{}"))
    orig_filename = f.filename
    ext = Path(orig_filename).suffix.lower()
    
    # 強制洗淨檔名為純英文
    safe_filename = f"task{ext}" if ext else "task.pdf"
    
    if "stem" not in cfg or not cfg["stem"]:
        cfg["stem"] = Path(orig_filename).stem

    file_data = f.read()

    # 🎯 PDF 記憶體精準裁剪邏輯
    page_range_str = cfg.get("page_range", "").strip()
    if ext == ".pdf" and page_range_str:
        try:
            import fitz  # PyMuPDF
            src_doc = fitz.open(stream=file_data, filetype="pdf")
            total_pages = src_doc.page_count
            
            target_pages = []
            for part in page_range_str.split(','):
                if '-' in part:
                    start, end = part.split('-')
                    s = max(0, int(start.strip()) - 1)
                    e = min(total_pages, int(end.strip()))
                    target_pages.extend(range(s, e))
                else:
                    if part.strip().isdigit():
                        idx = int(part.strip()) - 1
                        if 0 <= idx < total_pages:
                            target_pages.append(idx)
            
            if target_pages:
                target_pages = sorted(set(target_pages))
                src_doc.select(target_pages)   # 原地只保留所选页（insert_pdf 无 select 参数，旧写法会抛错回退整本）
                file_data = src_doc.tobytes()
                print(f"PDF cropped to range [{page_range_str}] -> {len(target_pages)} pages")
            src_doc.close()
        except Exception as e:
            print(f"[warn] PDF crop failed: {str(e)}")

    _prune_tasks()   # 清理过期的已完成任务，避免内存无限增长
    task_id = str(uuid.uuid4())
    _task_statuses[task_id] = {
        "state": "running", "msg": "準備上傳文件並初始化引擎...",
        "result": None, "n": 0, "ocr_ready": False, "ocr_pages": None, "trans": None,
        "t0": time.time(),
    }

    threading.Thread(
        target=_async_worker, args=(task_id, file_data, safe_filename, cfg), daemon=True
    ).start()

    return jsonify({"status": "submitted", "task_id": task_id})

@app.route("/api/task_status", methods=["GET"])
def task_status():
    tid = request.args.get("task_id")
    st = _task_statuses.get(tid)
    if not tid or st is None:
        return jsonify({"state": "error", "msg": "找不到指定的任务 ID"}), 404
    # 只回轻量进度；正文（OCR 全文 / 各页译文）经 /api/task_pages 增量获取，
    # 避免每 3 秒把整本书重传一遍。app.config 已护航 UTF-8，永不报 latin-1。
    trans = st.get("trans") or []
    trans_done = [i for i, t in enumerate(trans) if t is not None]
    return jsonify({
        "state":     st.get("state"),
        "msg":       st.get("msg", ""),
        "n":         st.get("n", 0),
        "ocr_ready": bool(st.get("ocr_ready")),
        "trans_done": trans_done,
    })

@app.route("/api/task_pages", methods=["GET"])
def task_pages():
    """按需取正文：kind=ocr 取全部 OCR 页；kind=trans&ids=1,2 取指定页译文；kind=result 取最终结果。"""
    tid = request.args.get("task_id")
    st = _task_statuses.get(tid)
    if not tid or st is None:
        return jsonify({"error": "找不到任务"}), 404
    kind = request.args.get("kind", "ocr")
    if kind == "ocr":
        return jsonify({"ocr_pages": st.get("ocr_pages") or []})
    if kind == "trans":
        trans = st.get("trans") or []
        out = {}
        for part in request.args.get("ids", "").split(","):
            part = part.strip()
            if part.isdigit():
                i = int(part)
                if 0 <= i < len(trans) and trans[i] is not None:
                    out[str(i)] = trans[i]
        return jsonify({"trans": out})
    if kind == "result":
        return jsonify(st.get("result") or {})
    return jsonify({"error": "未知 kind"}), 400

@app.route("/api/translate_page", methods=["POST"])
def translate_page_api():
    body = request.get_json(force=True)
    text = body.get("text", "").strip()
    cfg  = body.get("cfg", {})
    if not text: return {"ok": False, "error": "空文本"}, 400
    try:
        tr = _translate(text, cfg.get("target", "中文"), cfg)
        return {"ok": True, "translation": tr}
    except Exception as e:
        return {"ok": False, "error": str(e)}, 500

@app.route("/api/export_docx", methods=["POST"])
def export_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"error": "请先安装 python-docx"}, 500
        
    body   = request.get_json(force=True)
    pages  = body.get("pages", [])
    stem   = body.get("stem", "文献")
    source = body.get("source", "日语")
    target = body.get("target", "中文")
    
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = sec.right_margin = Cm(3)
        
    tp = doc.add_heading(f"{stem} {source}{target}对照", level=1)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    def _md_to_plain(text: str) -> str:
        text = _strip_md_noise(text)
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
        text = re.sub(r'_{1,2}(.*?)_{1,2}', r'\1', text)
        text = re.sub(r'^>\s?', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\|(.+)\|$', lambda m: '  '.join(c.strip() for c in m.group(1).split('|')), text, flags=re.MULTILINE)
        text = re.sub(r'^[-:|  ]+$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
        return text.strip()

    def _set_para_shading(p, fill='F0EBD8'):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        pPr.append(shd)

    def _set_cjk_font(run, font_name):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), 'Times New Roman')

    for pg in pages:
        pnum = pg.get("page_num", 1)
        ocr  = _md_to_plain(pg.get("ocr", ""))
        trans = _md_to_plain(pg.get("trans", ""))
        
        doc.add_heading(f"第 {pnum} 页", level=2)
        for blk in ocr.split('\n\n'):
            if not blk.strip(): continue
            p = doc.add_paragraph()
            _set_para_shading(p, 'F0EBD8')
            run = p.add_run(blk.strip())
            run.font.size = Pt(10)
            _set_cjk_font(run, '游明朝')
            p.paragraph_format.space_after = Pt(2)
            
        for blk in trans.split('\n\n'):
            if not blk.strip(): continue
            p = doc.add_paragraph()
            run = p.add_run(blk.strip())
            run.font.size = Pt(11)
            _set_cjk_font(run, '宋体')
            p.paragraph_format.space_after = Pt(4)

        # 校读笔记（存疑/想法/待查）—— 附在本页译文之后，浅蓝底楷体区分
        note = (pg.get("note") or "").strip()
        if note:
            lp = doc.add_paragraph()
            lr = lp.add_run("【筆記】")
            lr.bold = True
            lr.font.size = Pt(10)
            _set_cjk_font(lr, '楷体')
            lp.paragraph_format.space_after = Pt(1)
            for blk in note.split('\n\n'):
                blk = blk.strip()
                if not blk: continue
                p = doc.add_paragraph()
                _set_para_shading(p, 'E6EEF4')
                run = p.add_run(blk)
                run.font.size = Pt(10)
                run.italic = True
                _set_cjk_font(run, '楷体')
                p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe = re.sub(r'[^\w一-鿿぀-ヿ-]', '_', stem)[:40]
    return Response(
        buf.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe + '_对照.docx')}"},
    )

@app.route("/")
def index():
    return send_from_directory(Path(__file__).parent, "index.html")

@app.route("/api/health")
def health():
    """供本地启动器确认端口上运行的是本工具，而不是其他程序。"""
    return jsonify({"ok": True, "app": "ja-lit-local"})

if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 7860))
    # 本地版默认只监听本机，避免同一局域网里的其他设备访问页面和任务数据。
    # 确有局域网共享需求时，可显式设置 HOST=0.0.0.0。
    host = os.environ.get("HOST", "127.0.0.1")
    app.run(host=host, port=port, debug=False)
